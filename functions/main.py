# Menggunakan library Firebase, bukan uvicorn
from firebase_functions import https_fn
from firebase_admin import initialize_app
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import tempfile
from typing import Dict, Any
import base64
import io

initialize_app()

# Inisialisasi FastAPI
app = FastAPI()

# Tambahkan CORS Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Izinkan semua untuk development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Endpoint Anda sekarang berada di bawah 'app' FastAPI
@app.post("/api/export-word")
async def export_word(data: Dict[str, Any]):
    # Kode untuk membuat dokumen Word tetap sama persis
    doc = Document()

    # Judul
    title = doc.add_heading('Laporan Hasil Tes Moral Intelligence', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    participant_name = data.get('participantName', 'Tanpa Nama')
    doc.add_paragraph(f"Nama Peserta: {participant_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Tanggal Tes: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Ringkasan Hasil
    doc.add_heading('Ringkasan Hasil', level=2)
    overall_score = data.get('overallScore', 0)
    score_category = data.get('scoreCategory', '')
    score_para = doc.add_paragraph()
    score_para.add_run('Skor Keseluruhan: ').bold = True
    score_run = score_para.add_run(f"{overall_score}/100 ({score_category})")
    score_run.bold = True
    score_run.font.size = Pt(14)
    doc.add_paragraph(data.get('scoreInterpretation', ''))

    # Tabel Hasil Detail
    doc.add_heading('Rincian per Aspek', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aspek'
    hdr_cells[1].text = 'Skor'
    hdr_cells[2].text = 'Kategori'
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    aspects_data = data.get('aspects', {})
    aspect_categories = data.get('aspectCategories', {})
    for aspect, score in aspects_data.items():
        row_cells = table.add_row().cells
        display_name = aspect.replace("hatiNurani", "Hati Nurani").replace("pengendalianDiri", "Pengendalian Diri").replace("kebaikanHati", "Kebaikan Hati").capitalize()
        row_cells[0].text = display_name
        row_cells[1].text = str(round(score))
        row_cells[2].text = aspect_categories.get(aspect, '')
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Gambar Grafik Radar
    chart_image_b64 = data.get('chartImage')
    if chart_image_b64:
        image_data = base64.b64decode(chart_image_b64.split(',')[1])
        image_stream = io.BytesIO(image_data)
        doc.add_heading('Grafik Radar', level=2)
        doc.add_picture(image_stream, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Simpan ke file sementara
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return tmp.name

# Mendefinisikan function utama yang akan di-deploy
@https_fn.on_request()
def moral_intelligence_api(req: https_fn.Request) -> https_fn.Response:
    # Menghubungkan request ke aplikasi FastAPI
    return https_fn.Response.from_asgi(app, req)
    