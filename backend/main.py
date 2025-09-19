from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os
import tempfile
from typing import Dict, Any
import base64
import io

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def remove_file(path: str) -> None:
    try:
        os.unlink(path)
    except Exception as e:
        print(f"Error removing file {path}: {e}")

@app.post("/api/export-word")
async def export_word(data: Dict[str, Any], background_tasks: BackgroundTasks):
    try:
        doc = Document()
        
        # [FORMAT BARU] Mengatur font default untuk seluruh dokumen
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # [FORMAT BARU] Judul utama 14pt dan Bold
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('Laporan Hasil Tes Moral Intelligence')
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'

        # Info Peserta
        participant_name = data.get('participantName', 'Tanpa Nama')
        doc.add_paragraph(f"Nama Peserta: {participant_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Tanggal Tes: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Interpretasi akhir/keseluruhan (dimajukan ke atas)
        doc.add_paragraph().add_run('Interpretasi Keseluruhan').bold = True
        doc.add_paragraph(data.get('scoreInterpretation', ''))
        
        # Tabel Hasil Detail (sekarang dengan deskripsi)
        doc.add_paragraph().add_run('Skor dan Deskripsi per Aspek').bold = True
        table = doc.add_table(rows=1, cols=4) # [PERUBAHAN] Menjadi 4 kolom
        table.style = 'Table Grid'
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(0.7)
        table.columns[2].width = Inches(1.0)
        table.columns[3].width = Inches(3.0) # Kolom baru untuk deskripsi

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aspek'
        hdr_cells[1].text = 'Skor'
        hdr_cells[2].text = 'Kategori'
        hdr_cells[3].text = 'Deskripsi Singkat' # Header baru
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        aspects_data = data.get('aspects', {})
        aspect_categories = data.get('aspectCategories', {})
        aspect_interpretations = data.get('aspectInterpretations', {}) # Mengambil data deskripsi
        
        for aspect, score in aspects_data.items():
            row_cells = table.add_row().cells
            display_name = aspect.replace("hatiNurani", "Hati Nurani").replace("pengendalianDiri", "Pengendalian Diri").replace("kebaikanHati", "Kebaikan Hati").capitalize()
            
            row_cells[0].text = display_name
            row_cells[1].text = str(round(score))
            row_cells[2].text = aspect_categories.get(aspect, '')
            row_cells[3].text = aspect_interpretations.get(aspect, '') # Menambahkan deskripsi
            
            # Format sel
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Gambar Grafik Radar
        chart_image_b64 = data.get('chartImage')
        if chart_image_b64:
            try:
                image_data = base64.b64decode(chart_image_b64.split(',')[1])
                image_stream = io.BytesIO(image_data)
                doc.add_paragraph().add_run('Grafik Radar').bold = True
                doc.add_picture(image_stream, width=Inches(5.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as img_e:
                print(f"Could not process image: {img_e}")
                doc.add_paragraph("(Gagal memuat gambar grafik)")
        
        # Simpan ke file sementara dan kirim
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        background_tasks.add_task(remove_file, tmp_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return FileResponse(
            tmp_path,
            filename=f"Hasil_Tes_Moral_Intelligence_{timestamp}.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {e}")

# (Kode di bawah ini tidak perlu diubah)
@https_fn.on_request()
def moral_intelligence_api(req: https_fn.Request) -> https_fn.Response:
    return https_fn.Response.from_asgi(app, req)